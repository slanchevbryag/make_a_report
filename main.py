import coordinates
import weather
import docx

date = input("Укажите дату в формате yyyy-MM-dd: ")
filepath_track = input("Укажите путь к файлу gpx: ")

point = coordinates.getting_coordinats(filepath_track)
point_coordinates = point[0]
point_coordinates = " ".join(point_coordinates)

weather_and_astro = weather.weather_by_terrain(point_coordinates, date)

doc = docx.Document()

doc.add_heading('Техническое описание', 0)
doc.add_heading('1 день', 1)
doc.paragraphs[1].runs[0].add_break()

par2 = doc.add_paragraph('Дата: ')
doc.paragraphs[2].runs[0].bold = True
par2.add_run(date)

par3 = doc.add_paragraph('Восход: ')
doc.paragraphs[3].runs[0].bold = True
par3.add_run(weather_and_astro['sunrise'])
par3.add_run('  Закат: ').bold = True
par3.add_run(weather_and_astro['sunset'])

par4 = doc.add_paragraph('Max температура: ')
doc.paragraphs[4].runs[0].bold = True
par4.add_run(f'{weather_and_astro['maxtempC']} C')
par4.add_run('  Min температура: ').bold = True
par4.add_run(f'{weather_and_astro['mintempC']} C')

pat5 = doc.add_paragraph(weather_and_astro["weatherDesc"])

doc.save('report.docx')
