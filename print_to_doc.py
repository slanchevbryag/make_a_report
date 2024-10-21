def print_to_doc(weather_and_astro, date):
    """
    Эта функция создаёт документ с отчётом.
    """

    import docx
    from docx.shared import Mm
    import float_img

    doc = docx.Document()

    doc.add_heading("Техническое описание", 0)
    doc.add_heading("1 день", 1)
    doc.paragraphs[1].runs[0].add_break()

    try:
        map = doc.add_paragraph()
        float_img.add_float_picture(
            map, "map.png", width=Mm(70), pos_x=Mm(30), pos_y=Mm(60)
        )
    except FileNotFoundError:
        print('Файл не найден')
        map = doc.add_paragraph()
        float_img.add_float_picture(
            map, "dead_smiley.png", width=Mm(70), pos_x=Mm(30), pos_y=Mm(60)
        )

    par3 = doc.add_paragraph("  Дата: ")
    doc.paragraphs[3].runs[0].bold = True
    par3.add_run(date)

    par4 = doc.add_paragraph("  Восход: ")
    doc.paragraphs[4].runs[0].bold = True
    par4.add_run(weather_and_astro["sunrise"])
    par4.add_run("  Закат: ").bold = True
    par4.add_run(weather_and_astro["sunset"])

    par5 = doc.add_paragraph("  Max температура: ")
    doc.paragraphs[5].runs[0].bold = True
    par5.add_run(f'{weather_and_astro['maxtempC']} C')

    par6 = doc.add_paragraph("  Min температура: ")
    doc.paragraphs[6].runs[0].bold = True
    par6.add_run(f'{weather_and_astro['mintempC']} C')

    par7 = doc.add_paragraph()
    par7.add_run(f'  {weather_and_astro['weatherDesc']}')

    doc.save("report.docx")


if __name__ in "__main__":
    print_to_doc(
        weather_and_astro={
            "sunrise": "04:01 AM",
            "sunset": "08:59 PM",
            "moonrise": "06:12 AM",
            "moonset": "10:42 PM",
            "moon_phase": "Waxing Crescent",
            "moon_illumination": "10",
            "maxtempC": "26",
            "mintempC": "14",
            "weatherDesc": "Солнечно",
        },
        date="2022-05-20",
    )
