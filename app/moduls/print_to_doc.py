import os
from typing import Any

from app.utilities import float_img

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm


def print_weather_and_trake(
        weather_and_astro: dict[str, Any], date: str,
        img_width: int, img_length: int, distance: float, days_travels: int) -> None:
    """
    Эта функция создаёт документ с отчётом и заносит в него миниатюру карты
    с треком и астрономические и погодные данные.
    """

    doc = docx.Document("report.docx")

    base_pos_y = 60
    if days_travels != 0:
        base_pos_y = 40

    doc.add_heading(f"{days_travels+1} день", 1)

    try:
        mini_map = doc.add_paragraph()

        if img_length > img_width:

            float_img.add_float_picture(
                mini_map, "map.png", height=Mm(80), pos_x=Mm(30), pos_y=Mm(base_pos_y),
            )

        else:

            float_img.add_float_picture(
                mini_map, "map.png", width=Mm(80), pos_x=Mm(30), pos_y=Mm(base_pos_y),
            )

    except FileNotFoundError:
        print('Файл не найден')
        mini_map = doc.add_paragraph()
        float_img.add_float_picture(
            mini_map, "dead_smiley.png", height=Mm(80), pos_x=Mm(30), pos_y=Mm(base_pos_y),
        )

    par3 = doc.add_paragraph()
    par3.add_run("  Дата: ").bold = True
    par3.add_run(date)

    par4 = doc.add_paragraph()
    par4.add_run("  Восход: ").bold = True
    par4.add_run(weather_and_astro["sunrise"])
    par4.add_run("  Закат: ").bold = True
    par4.add_run(weather_and_astro["sunset"])

    par5 = doc.add_paragraph()
    par5.add_run("  Max температура: ").bold = True
    par5.add_run(f'{weather_and_astro['maxtempC']} C')

    par6 = doc.add_paragraph()
    par6.add_run("  Min температура: ").bold = True
    par6.add_run(f'{weather_and_astro['mintempC']} C')

    par7 = doc.add_paragraph()
    par7.add_run(f'  {weather_and_astro['weatherDesc']}')

    par8 = doc.add_paragraph()
    par8.add_run("  Пройдено: ").bold = True
    par8.add_run(f'{distance} км')

    doc.save("report.docx")


def print_travel_notes(img_width: int, img_length: int) -> None:
    '''
    Эта функция добавляет в файл отчёта путевые заметки.
    '''

    with open('draft.txt', 'r', encoding='utf-8') as d:
        draft = d.read()

    doc = docx.Document("report.docx")

    if img_length > img_width:

        for _ in range(1, 3):
            _ = doc.add_paragraph()

    par_draft = doc.add_paragraph()
    par_draft.add_run(draft)

    doc.save("report.docx")


def print_foto(temp_path: str = "temp") -> None:
    '''
    Эта функция добавляет отобранные пользователем фото в отчёт.
    '''

    doc = docx.Document("report.docx")

    foto_in_temp = []
    try:

        for file_in_dir in os.listdir(path=temp_path):
            foto_in_temp.append(file_in_dir)

    except FileNotFoundError:
        print('Файлы не найдены')

    foto_in_temp.sort()

    try:
        for num_foto in range(0, len(foto_in_temp), 2):
            par_foto = doc.add_paragraph()
            par_foto.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run1_foto = par_foto.add_run()
            run1_foto.add_picture(os.path.join('temp', foto_in_temp[num_foto]), width=Mm(69))

            run2_foto = par_foto.add_run()
            run2_foto.add_picture(os.path.join("temp", foto_in_temp[num_foto+1]), width=Mm(69))
    except IndexError:
        pass

    doc.add_page_break()

    doc.save("report.docx")
